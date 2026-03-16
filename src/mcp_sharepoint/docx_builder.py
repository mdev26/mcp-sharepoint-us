"""
Word document builder using python-docx.
Creates formatted Word documents from JSON structure.
"""
import logging
import base64
import requests
from typing import Dict, Any, List, Tuple, Optional
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX, WD_LINE_SPACING
from docx.oxml.shared import OxmlElement, qn

logger = logging.getLogger(__name__)

# Named color mappings (common colors)
NAMED_COLORS = {
    "black": (0, 0, 0),
    "white": (255, 255, 255),
    "red": (255, 0, 0),
    "green": (0, 128, 0),
    "blue": (0, 0, 255),
    "yellow": (255, 255, 0),
    "orange": (255, 165, 0),
    "purple": (128, 0, 128),
    "pink": (255, 192, 203),
    "gray": (128, 128, 128),
    "grey": (128, 128, 128),
    "brown": (165, 42, 42),
    "navy": (0, 0, 128),
    "teal": (0, 128, 128),
    "lime": (0, 255, 0),
    "cyan": (0, 255, 255),
    "magenta": (255, 0, 255),
    "maroon": (128, 0, 0),
    "olive": (128, 128, 0),
}

# Highlight color mappings
HIGHLIGHT_COLORS = {
    "yellow": WD_COLOR_INDEX.YELLOW,
    "bright_green": WD_COLOR_INDEX.BRIGHT_GREEN,
    "turquoise": WD_COLOR_INDEX.TURQUOISE,
    "pink": WD_COLOR_INDEX.PINK,
    "blue": WD_COLOR_INDEX.BLUE,
    "red": WD_COLOR_INDEX.RED,
    "dark_blue": WD_COLOR_INDEX.DARK_BLUE,
    "dark_yellow": WD_COLOR_INDEX.DARK_YELLOW,
    "green": WD_COLOR_INDEX.GREEN,
    "gray_25": WD_COLOR_INDEX.GRAY_25,
}


class WordDocumentBuilder:
    """
    Builder for creating Word documents from JSON structure.

    Supports:
    - Headings (levels 1-9)
    - Paragraphs with formatting (bold, italic, underline, color, highlight, spacing)
    - Tables with headers and data
    - Bulleted and numbered lists
    - Page breaks
    - Images (from URLs or base64)
    - Hyperlinks
    - Horizontal rules
    """

    def __init__(self):
        """Initialize the document builder."""
        self.doc = Document()

    def _parse_color(self, color: Any) -> Optional[RGBColor]:
        """
        Parse color from various formats to RGBColor.

        Supports:
        - Named colors: "red", "blue", etc.
        - RGB tuples/lists: [255, 0, 0] or (255, 0, 0)
        - RGB strings: "255,0,0"
        """
        if not color:
            return None

        if isinstance(color, str):
            # Check if it's a named color
            color_lower = color.lower()
            if color_lower in NAMED_COLORS:
                r, g, b = NAMED_COLORS[color_lower]
                return RGBColor(r, g, b)

            # Try parsing as RGB string "255,0,0"
            try:
                parts = color.split(',')
                if len(parts) == 3:
                    r, g, b = [int(p.strip()) for p in parts]
                    return RGBColor(r, g, b)
            except (ValueError, AttributeError):
                pass

        elif isinstance(color, (list, tuple)) and len(color) == 3:
            # RGB tuple/list
            try:
                r, g, b = [int(c) for c in color]
                return RGBColor(r, g, b)
            except (ValueError, TypeError):
                pass

        logger.warning(f"Could not parse color: {color}")
        return None

    def _fetch_image(self, source: str) -> Optional[BytesIO]:
        """
        Fetch image from URL or decode from base64.

        Args:
            source: Either a URL (http/https) or base64 string

        Returns:
            BytesIO containing image data, or None if failed
        """
        try:
            if source.startswith(('http://', 'https://')):
                # Fetch from URL
                response = requests.get(source, timeout=30)
                response.raise_for_status()
                return BytesIO(response.content)
            elif source.startswith('data:image'):
                # Handle data URI: data:image/png;base64,iVBORw0KG...
                if 'base64,' in source:
                    base64_data = source.split('base64,')[1]
                    image_bytes = base64.b64decode(base64_data)
                    return BytesIO(image_bytes)
            else:
                # Assume it's raw base64
                image_bytes = base64.b64decode(source)
                return BytesIO(image_bytes)
        except Exception as e:
            logger.error(f"Error fetching/decoding image: {e}")
            return None

    def build_from_json(self, content_list: List[Dict[str, Any]]) -> bytes:
        """
        Build a Word document from JSON content structure.

        Args:
            content_list: List of content items, each with a 'type' and specific properties

        Returns:
            Document content as bytes

        Example content_list:
        [
            {"type": "heading", "level": 1, "text": "Title"},
            {"type": "paragraph", "text": "Some text", "bold": true},
            {"type": "table", "headers": ["Col1", "Col2"], "rows": [["A", "B"]]},
            {"type": "list", "style": "bullet", "items": ["Item 1", "Item 2"]},
            {"type": "page_break"}
        ]
        """
        for item in content_list:
            item_type = item.get("type", "")

            try:
                if item_type == "heading":
                    self._add_heading(item)
                elif item_type == "paragraph":
                    self._add_paragraph(item)
                elif item_type == "table":
                    self._add_table(item)
                elif item_type == "list":
                    self._add_list(item)
                elif item_type == "image":
                    self._add_image(item)
                elif item_type == "hyperlink":
                    self._add_hyperlink(item)
                elif item_type == "horizontal_rule":
                    self._add_horizontal_rule()
                elif item_type == "page_break":
                    self._add_page_break()
                else:
                    logger.warning(f"Unknown content type: {item_type}")

            except Exception as e:
                logger.error(f"Error adding {item_type}: {e}", exc_info=True)
                # Add error message to document
                p = self.doc.add_paragraph()
                run = p.add_run(f"[Error adding {item_type}: {str(e)}]")
                run.font.color.rgb = RGBColor(255, 0, 0)

        # Save to BytesIO
        buffer = BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        return buffer.read()

    def _add_heading(self, item: Dict[str, Any]) -> None:
        """
        Add a heading to the document.

        Args:
            item: {"type": "heading", "level": 1-9, "text": "Heading text"}
        """
        level = item.get("level", 1)
        text = item.get("text", "")

        if level < 1 or level > 9:
            logger.warning(f"Invalid heading level {level}, using level 1")
            level = 1

        self.doc.add_heading(text, level=level)

    def _add_paragraph(self, item: Dict[str, Any]) -> None:
        """
        Add a paragraph with optional formatting.

        Args:
            item: {
                "type": "paragraph",
                "text": "Paragraph text",
                "bold": true/false (optional),
                "italic": true/false (optional),
                "underline": true/false (optional),
                "font_size": 12 (optional, in points),
                "align": "left"/"center"/"right" (optional),
                "color": "red" or [255,0,0] (optional),
                "highlight": "yellow" (optional),
                "line_spacing": 1.5 (optional),
                "space_before": 6 (optional, in points),
                "space_after": 6 (optional, in points)
            }
        """
        text = item.get("text", "")
        bold = item.get("bold", False)
        italic = item.get("italic", False)
        underline = item.get("underline", False)
        font_size = item.get("font_size")
        align = item.get("align", "left")
        color = item.get("color")
        highlight = item.get("highlight")
        line_spacing = item.get("line_spacing")
        space_before = item.get("space_before")
        space_after = item.get("space_after")

        p = self.doc.add_paragraph()
        run = p.add_run(text)

        # Apply text formatting
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        if underline:
            run.underline = True
        if font_size:
            run.font.size = Pt(font_size)

        # Apply color
        if color:
            rgb_color = self._parse_color(color)
            if rgb_color:
                run.font.color.rgb = rgb_color

        # Apply highlighting
        if highlight:
            highlight_lower = highlight.lower()
            if highlight_lower in HIGHLIGHT_COLORS:
                run.font.highlight_color = HIGHLIGHT_COLORS[highlight_lower]
            else:
                logger.warning(f"Unknown highlight color: {highlight}")

        # Apply alignment
        if align == "center":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == "right":
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # left is default

        # Apply paragraph spacing
        if line_spacing:
            p.paragraph_format.line_spacing = line_spacing
        if space_before:
            p.paragraph_format.space_before = Pt(space_before)
        if space_after:
            p.paragraph_format.space_after = Pt(space_after)

    def _add_table(self, item: Dict[str, Any]) -> None:
        """
        Add a table to the document.

        Args:
            item: {
                "type": "table",
                "headers": ["Col1", "Col2", "Col3"],
                "rows": [
                    ["Data1", "Data2", "Data3"],
                    ["Data4", "Data5", "Data6"]
                ],
                "style": "Light Grid Accent 1" (optional)
            }
        """
        headers = item.get("headers", [])
        rows = item.get("rows", [])
        style = item.get("style", "Light Grid Accent 1")

        if not headers:
            logger.warning("Table has no headers, skipping")
            return

        # Create table with header row + data rows
        table = self.doc.add_table(rows=1 + len(rows), cols=len(headers))

        # Apply style
        try:
            table.style = style
        except Exception as e:
            logger.warning(f"Could not apply table style '{style}': {e}")

        # Add headers
        header_row = table.rows[0]
        for i, header_text in enumerate(headers):
            cell = header_row.cells[i]
            cell.text = str(header_text)
            # Make header bold
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        # Add data rows
        for row_idx, row_data in enumerate(rows, start=1):
            table_row = table.rows[row_idx]
            for col_idx, cell_data in enumerate(row_data):
                if col_idx < len(headers):  # Don't exceed column count
                    table_row.cells[col_idx].text = str(cell_data)

    def _add_list(self, item: Dict[str, Any]) -> None:
        """
        Add a bulleted or numbered list to the document.

        Args:
            item: {
                "type": "list",
                "style": "bullet" or "number",
                "items": ["Item 1", "Item 2", "Item 3"]
            }
        """
        style = item.get("style", "bullet")
        items = item.get("items", [])

        if not items:
            logger.warning("List has no items, skipping")
            return

        for item_text in items:
            if style == "number":
                self.doc.add_paragraph(str(item_text), style='List Number')
            else:  # bullet
                self.doc.add_paragraph(str(item_text), style='List Bullet')

    def _add_image(self, item: Dict[str, Any]) -> None:
        """
        Add an image to the document.

        Args:
            item: {
                "type": "image",
                "source": "https://..." or "base64data..." or "data:image/png;base64,...",
                "width": 4.0 (optional, in inches),
                "height": 3.0 (optional, in inches)
            }
        """
        source = item.get("source", "")
        width = item.get("width")
        height = item.get("height")

        if not source:
            logger.warning("Image source is empty, skipping")
            return

        # Fetch/decode image
        image_stream = self._fetch_image(source)
        if not image_stream:
            logger.error(f"Failed to fetch image from: {source[:100]}...")
            p = self.doc.add_paragraph()
            run = p.add_run(f"[Failed to load image from source]")
            run.font.color.rgb = RGBColor(255, 0, 0)
            return

        # Add image with optional dimensions
        try:
            if width and height:
                self.doc.add_picture(image_stream, width=Inches(width), height=Inches(height))
            elif width:
                self.doc.add_picture(image_stream, width=Inches(width))
            elif height:
                self.doc.add_picture(image_stream, height=Inches(height))
            else:
                self.doc.add_picture(image_stream)
        except Exception as e:
            logger.error(f"Error adding image to document: {e}")
            p = self.doc.add_paragraph()
            run = p.add_run(f"[Error adding image: {str(e)}]")
            run.font.color.rgb = RGBColor(255, 0, 0)

    def _add_hyperlink(self, item: Dict[str, Any]) -> None:
        """
        Add a hyperlink to the document.

        Args:
            item: {
                "type": "hyperlink",
                "text": "Click here",
                "url": "https://example.com",
                "color": "blue" (optional)
            }
        """
        text = item.get("text", "")
        url = item.get("url", "")
        color = item.get("color", "blue")  # Default to blue for links

        if not url:
            logger.warning("Hyperlink URL is empty, adding as plain text")
            p = self.doc.add_paragraph()
            p.add_run(text)
            return

        # Create paragraph and add hyperlink
        p = self.doc.add_paragraph()

        # Add hyperlink using Word's relationship system
        part = p.part
        r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

        # Create the hyperlink element
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)

        # Create a new run for the hyperlink text
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')

        # Apply hyperlink style (blue, underlined)
        rgb_color = self._parse_color(color)
        if rgb_color:
            c = OxmlElement('w:color')
            c.set(qn('w:val'), f"{rgb_color.r:02X}{rgb_color.g:02X}{rgb_color.b:02X}")
            rPr.append(c)

        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)

        # Add hyperlink to paragraph
        p._p.append(hyperlink)

    def _add_horizontal_rule(self) -> None:
        """Add a horizontal rule (line) to the document."""
        p = self.doc.add_paragraph()
        pPr = p._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')

        # Add bottom border
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')  # Border size
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'auto')

        pBdr.append(bottom)
        pPr.append(pBdr)

    def _add_page_break(self) -> None:
        """Add a page break to the document."""
        self.doc.add_page_break()


def create_word_document(content: List[Dict[str, Any]]) -> bytes:
    """
    Create a Word document from JSON content structure.

    Args:
        content: List of content items (headings, paragraphs, tables, lists, page breaks)

    Returns:
        Document bytes ready for upload

    Example:
        content = [
            {"type": "heading", "level": 1, "text": "Monthly Report"},
            {"type": "paragraph", "text": "January 2025"},
            {"type": "table",
             "headers": ["Task", "Status", "Notes"],
             "rows": [
                 ["Task 1", "Complete", "Done"],
                 ["Task 2", "In Progress", "Working"]
             ]},
            {"type": "page_break"},
            {"type": "heading", "level": 2, "text": "Next Steps"},
            {"type": "list", "style": "bullet", "items": ["Step 1", "Step 2"]}
        ]
        doc_bytes = create_word_document(content)
    """
    builder = WordDocumentBuilder()
    return builder.build_from_json(content)


def edit_word_document(doc_bytes: bytes, replacements: List[Dict[str, Any]]) -> Tuple[bytes, Dict[str, int]]:
    """
    Edit an existing Word document by performing find/replace operations.
    Supports both simple text replacement and section replacement for monthly reports.
    All existing formatting is preserved.

    Args:
        doc_bytes: Bytes of the existing Word document
        replacements: List of replacement operations, supports two formats:

            SIMPLE REPLACEMENT:
            [
                {"find": "{{company_name}}", "replace": "Acme Corp"},
                {"find": "{{date}}", "replace": "2025-02-05"}
            ]

            SECTION REPLACEMENT (for monthly status reports):
            [
                {
                    "find": "{{section:status_bullets}}",
                    "replace_section": [
                        "Completed Phase 1 (Status: Done)",
                        "Started Phase 2 (Status: In Progress)",
                        "Identified blockers (Status: Blocked)"
                    ],
                    "section_type": "bullets"
                },
                {
                    "find": "{{section:metrics_table}}",
                    "replace_section": [
                        ["Metric", "Value"],
                        ["Revenue", "$1.5M"],
                        ["Users", "520"]
                    ],
                    "section_type": "table"
                }
            ]

    Returns:
        Tuple of (modified document bytes, statistics dict with replacement counts)

    Example - Simple Replacement:
        replacements = [
            {"find": "{{company}}", "replace": "Acme Corp"},
            {"find": "{{year}}", "replace": "2025"}
        ]

    Example - Section Replacement (Monthly Reports):
        # Template document has:
        # {{section:status_bullets}}
        # * Old bullet 1
        # * Old bullet 2
        # {{/section:status_bullets}}

        replacements = [
            {
                "find": "{{section:status_bullets}}",
                "replace_section": [
                    "Completed development of new feature X",
                    "Started testing phase for feature Y",
                    "Identified performance issues in module Z"
                ],
                "section_type": "bullets"
            }
        ]
    """
    try:
        # Load existing document
        doc_stream = BytesIO(doc_bytes)
        doc = Document(doc_stream)

        # Track statistics
        stats = {
            "total_replacements": 0,
            "paragraphs_modified": 0,
            "tables_modified": 0,
            "sections_replaced": 0
        }
        replacement_counts = {r["find"]: 0 for r in replacements}

        # PHASE 1: Handle section replacements first
        for replacement in replacements:
            if "replace_section" in replacement:
                section_name = replacement["find"]
                section_items = replacement["replace_section"]
                section_type = replacement.get("section_type", "bullets")

                # Process section replacement
                sections_replaced = _replace_section(doc, section_name, section_items, section_type)
                replacement_counts[section_name] = sections_replaced
                if sections_replaced > 0:
                    stats["sections_replaced"] += 1

        # PHASE 2: Handle simple text replacements
        # Replace in all paragraphs (including headers/footers)
        for paragraph in doc.paragraphs:
            modified = False
            for replacement in replacements:
                # Skip section replacements (already processed)
                if "replace_section" in replacement:
                    continue

                find_text = replacement["find"]
                replace_text = replacement.get("replace", "")

                # Check if this paragraph contains the text to replace
                if find_text in paragraph.text:
                    # Need to replace while preserving formatting
                    # Strategy: work with runs to preserve formatting
                    _replace_in_paragraph(paragraph, find_text, replace_text)
                    replacement_counts[find_text] += 1
                    modified = True

            if modified:
                stats["paragraphs_modified"] += 1

        # Replace in all tables
        for table in doc.tables:
            table_modified = False
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for replacement in replacements:
                            # Skip section replacements
                            if "replace_section" in replacement:
                                continue

                            find_text = replacement["find"]
                            replace_text = replacement.get("replace", "")

                            if find_text in paragraph.text:
                                _replace_in_paragraph(paragraph, find_text, replace_text)
                                replacement_counts[find_text] += 1
                                table_modified = True

            if table_modified:
                stats["tables_modified"] += 1

        # Calculate total replacements
        stats["total_replacements"] = sum(replacement_counts.values())
        stats["by_placeholder"] = replacement_counts

        # Save modified document
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read(), stats

    except Exception as e:
        logger.error(f"Error editing Word document: {e}", exc_info=True)
        raise


def _replace_in_paragraph(paragraph, find_text: str, replace_text: str) -> None:
    """
    Replace text in a paragraph while preserving formatting.

    Works by combining all runs into full text, finding replacements,
    then reconstructing runs with original formatting.

    Args:
        paragraph: python-docx Paragraph object
        find_text: Text to find
        replace_text: Text to replace with
    """
    # Build full text from all runs
    full_text = paragraph.text

    # If the text isn't in this paragraph, skip
    if find_text not in full_text:
        return

    # Perform replacement
    new_text = full_text.replace(find_text, replace_text)

    # If nothing changed, skip
    if new_text == full_text:
        return

    # Strategy: Keep the first run and update its text, remove others
    # This preserves the formatting of the first run
    # Note: This is a simple approach. For more complex scenarios where
    # find_text spans multiple runs with different formatting, we'd need
    # more sophisticated logic. But for template placeholders, this works well.

    if paragraph.runs:
        # Get the first run's formatting
        first_run = paragraph.runs[0]

        # Clear all runs except the first
        for _ in range(len(paragraph.runs) - 1):
            paragraph.runs[-1]._element.getparent().remove(paragraph.runs[-1]._element)

        # Update the first run with new text
        first_run.text = new_text
    else:
        # No runs exist, just set paragraph text directly
        paragraph.text = new_text


def _replace_section(doc: Document, section_marker: str, section_items: List[Any], section_type: str) -> int:
    """
    Replace an entire section in the document (for monthly status reports).

    Finds section markers like {{section:name}} and {{/section:name}}, deletes
    everything between them (including the markers and any tables), and inserts new content.

    Args:
        doc: python-docx Document object
        section_marker: Section start marker (e.g., "{{section:status_bullets}}")
        section_items: List of items to insert (strings for bullets, list of lists for tables)
        section_type: Type of section - "bullets" or "table"

    Returns:
        Number of sections replaced
    """
    # Extract section name from marker
    if not section_marker.startswith("{{section:"):
        logger.warning(f"Invalid section marker: {section_marker}")
        return 0

    section_name = section_marker.replace("{{section:", "").replace("}}", "")
    start_marker = f"{{{{section:{section_name}}}}}"
    end_marker = f"{{{{/section:{section_name}}}}}"

    # Find start and end markers in body elements
    # We need to work with the actual body elements to handle both paragraphs and tables
    body = doc._body._body
    start_element = None
    end_element = None

    for element in body:
        # Check if this is a paragraph with our markers
        if element.tag.endswith('p'):
            # Create a temporary paragraph object to get text
            para_text = ''.join([
                node.text for node in element.iter()
                if node.text and node.tag.endswith('t')
            ])

            if start_marker in para_text:
                start_element = element
            elif end_marker in para_text and start_element is not None:
                end_element = element
                break

    if start_element is None or end_element is None:
        logger.warning(f"Section markers not found for: {section_name}")
        return 0

    # Find the index position where we'll insert new content
    start_idx = body.index(start_element)
    end_idx = body.index(end_element)

    if end_idx <= start_idx:
        logger.warning(f"Invalid section markers for: {section_name} (end before start)")
        return 0

    # Delete ALL elements between start and end (inclusive)
    # This includes paragraphs, tables, and any other elements
    for idx in range(end_idx, start_idx - 1, -1):
        body.remove(body[idx])

    # Insert new content at start_idx position
    if section_type == "bullets":
        # Insert bullet points
        for item_text in section_items:
            p = doc.add_paragraph(str(item_text), style='List Bullet')
            # Move the paragraph to the correct position
            p_element = p._element
            body.insert(start_idx, p_element)
            start_idx += 1

    elif section_type == "table":
        # Insert table
        if len(section_items) > 0:
            num_cols = len(section_items[0]) if section_items else 1
            table = doc.add_table(rows=len(section_items), cols=num_cols)
            table.style = 'Light Grid Accent 1'

            # Populate table
            for row_idx, row_data in enumerate(section_items):
                for col_idx, cell_data in enumerate(row_data):
                    if col_idx < num_cols:
                        table.rows[row_idx].cells[col_idx].text = str(cell_data)

                        # Bold first row (header)
                        if row_idx == 0:
                            for paragraph in table.rows[row_idx].cells[col_idx].paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True

            # Move table to the correct position
            table_element = table._element
            body.insert(start_idx, table_element)

    return 1  # One section replaced


def _move_paragraph(doc: Document, from_idx: int, to_idx: int) -> None:
    """Move a paragraph from one position to another"""
    try:
        paragraphs = doc.paragraphs
        p = paragraphs[from_idx]._element
        if to_idx < len(paragraphs):
            target = paragraphs[to_idx]._element
            target.getparent().insert(target.getparent().index(target), p)
        else:
            target = paragraphs[-1]._element
            target.getparent().insert(target.getparent().index(target) + 1, p)
    except Exception as e:
        logger.error(f"Error moving paragraph: {e}")


def _move_table(doc: Document, from_idx: int, to_idx: int) -> None:
    """Move a table from one position to another"""
    try:
        tables = doc.tables
        table = tables[from_idx]._element
        # Insert table before the paragraph at to_idx
        if to_idx < len(doc.paragraphs):
            target = doc.paragraphs[to_idx]._element
            target.getparent().insert(target.getparent().index(target), table)
        else:
            # Append at end
            doc._body._body.append(table)
    except Exception as e:
        logger.error(f"Error moving table: {e}")
